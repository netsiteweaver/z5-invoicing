from __future__ import annotations

from datetime import datetime
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title_page(document: Document) -> None:
    title = document.add_heading("Point of Sale (POS) System — Proposal & Requirements", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_lines = [
        "Client: TBD",
        f"Date: {datetime.now().strftime('%Y-%m-%d')}",
        "Version: 1.0.0",
    ]
    for line in meta_lines:
        paragraph = document.add_paragraph(line)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraphs(document: Document, lines: List[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(6)


def add_bullets(document: Document, items: List[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.size = Pt(11)


def add_numbered(document: Document, items: List[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.size = Pt(11)


def add_table(document: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def section_overview(document: Document) -> None:
    add_heading(document, "1. Overview & Objectives")
    add_paragraphs(
        document,
        [
            "This document defines the scope, functional and non-functional requirements, security posture, and recommended technology stack for a new Point of Sale (POS) system.",
            "The primary product line is footwear, with support for selling additional non-footwear items. The solution must provide two distinct interfaces: a tablet-optimized Sales interface for in-store transactions and a Back Office interface for catalog, inventory, users, and reporting.",
            "Objectives include a fast, intuitive checkout experience, robust offline capability, accurate inventory tracking, and secure payment processing compliant with industry standards.",
        ],
    )


def section_scope(document: Document) -> None:
    add_heading(document, "2. Scope")
    document.add_paragraph("In scope:")
    add_bullets(
        document,
        [
            "Tablet-optimized POS (Sales) interface with offline-first capability",
            "Web-based Back Office for product, inventory, pricing, promotions, users, and reports",
            "Footwear-specific product variants (size, color) and general merchandise",
            "Barcode scanning, receipt printing, and cash drawer support",
            "Payment terminal integration (e.g., Stripe Terminal) with tokenization",
            "Multi-store locations and registers",
            "Real-time and background synchronization across devices",
        ],
    )
    document.add_paragraph("Out of scope (Phase 1):")
    add_bullets(
        document,
        [
            "Full e-commerce website (only optional future integration hooks)",
            "Advanced warehouse management (basic stock transfers only)",
            "Accounting system replacement (provide export/integration instead)",
        ],
    )


def section_stakeholders(document: Document) -> None:
    add_heading(document, "3. Stakeholders & User Roles")
    add_paragraphs(
        document,
        [
            "Stakeholders include Store Associates, Store Managers, Inventory Managers, Finance/Accounting, Operations, and Administrators.",
            "At minimum, the system will support role-based access controls (RBAC) with the roles summarized below.",
        ],
    )
    add_table(
        document,
        headers=["Role", "Responsibilities", "Permissions (summary)"],
        rows=[
            [
                "Sales Associate",
                "Process sales, returns/exchanges, customer lookup, issue receipts",
                "Checkout, view catalog, apply allowed discounts, hold/resume",
            ],
            [
                "Store Manager",
                "Override discounts, manage users in store, daily reconciliation",
                "All associate permissions + overrides, end-of-day reports",
            ],
            [
                "Inventory Manager",
                "Maintain products, variants, stock counts, transfers",
                "Create/edit products, adjust inventory, receive stock",
            ],
            [
                "Admin",
                "System configuration, tax rules, multi-store settings, audits",
                "Full back-office access, security settings, audit exports",
            ],
        ],
    )


def section_interfaces(document: Document) -> None:
    add_heading(document, "4. Interfaces")
    document.add_paragraph("Sales (Tablet POS):")
    add_bullets(
        document,
        [
            "Responsive, touch-friendly UI optimized for 10–13" + "\u2033" + " tablets",
            "Scan/search products, build cart, apply discounts, collect payment",
            "Customer lookup/enrollment, receipt via print/email/SMS",
            "Offline-first with background synchronization",
        ],
    )
    document.add_paragraph("Back Office (Web Admin):")
    add_bullets(
        document,
        [
            "Product, category, variant management (sizes, colors)",
            "Pricing, promotions, and tax rules",
            "Inventory, stock adjustments, transfers, and purchase orders",
            "User and role management, store and register configuration",
            "Operational and financial reporting",
        ],
    )


def section_functional_requirements(document: Document) -> None:
    add_heading(document, "5. Functional Requirements")

    document.add_paragraph("5.1 Sales (Tablet POS)")
    add_bullets(
        document,
        [
            "Secure sign-in with role-based access",
            "Product scan (barcode), quick search, and variant selection (size, color)",
            "Cart operations: add/remove items, quantity edit, notes",
            "Discounts: per-line and cart-level (role/approval-gated)",
            "Tax calculation by store/location rules",
            "Payments: card (terminal), cash, gift card, split payments",
            "Refunds and exchanges with receipt/lookup",
            "Receipt printing/email/SMS; reprint on demand",
            "Hold/resume transactions; parked sales",
            "Customer profiles, simple loyalty (points/balance), and marketing opt-in",
            "Inventory lookup by store; low-stock warnings",
            "Shift management: open/close register, cash drawer operations",
            "Offline mode: queue transactions and sync on reconnect",
        ],
    )

    document.add_paragraph("5.2 Back Office (Web Admin)")
    add_bullets(
        document,
        [
            "Product catalog: categories, variants (size/color), SKUs, barcodes",
            "Pricing and promotions: price lists, promo codes, time-bound discounts",
            "Tax configuration per jurisdiction; rounding rules",
            "Inventory: stock levels, cycle counts, adjustments, transfers",
            "Suppliers and purchase orders: receiving and reconciliation",
            "Users, roles, and permissions; audit trails",
            "Stores and registers configuration; receipt templates",
            "Reports: sales, margins, tax, inventory movement, cashier performance",
            "Data import/export (CSV) for products and inventory",
            "Multi-store management with per-store overrides",
        ],
    )

    document.add_paragraph("5.3 Cross-Cutting & Integrations")
    add_bullets(
        document,
        [
            "Payment terminal integration (e.g., Stripe Terminal, Adyen, or Square)",
            "Receipt printers (ESC/POS over network) and cash drawers",
            "Barcode scanners (BLE/USB HID) and handheld scanners",
            "Background sync and conflict resolution (last-write-wins + merge rules)",
            "Webhooks/ETL for accounting systems (e.g., export to CSV/S3)",
            "Localization (currency, date/number formats); multi-currency display",
            "Accessibility (WCAG 2.1 AA) for back-office; large-touch targets for POS",
        ],
    )


def section_non_functional(document: Document) -> None:
    add_heading(document, "6. Non-Functional Requirements & SLAs")
    add_table(
        document,
        headers=["Category", "Target/Expectation"],
        rows=[
            ["Performance (POS UI)", "Initial load ≤ 3s on modern tablets; interactions ≤ 100ms"],
            ["Transaction time", "Add item ≤ 300ms; tender to receipt ≤ 5s (online)"],
            ["Availability", ">= 99.9% monthly for cloud services"],
            ["Offline tolerance", "Operate offline ≥ 72 hours; seamless resync on reconnect"],
            ["Data durability", ">= 11x9s for production databases (managed service)"],
            ["Scalability", "Support 100+ concurrent registers across stores"],
            ["Security", "TLS 1.2+ in transit, AES-256 at rest; RBAC; audit logs"],
            ["Usability", "Optimized touch UI; ADA/WCAG 2.1 AA (back office)"],
            ["Observability", "Structured logs, metrics, traces; alerting on SLOs"],
            ["Backup & DR", "Daily snapshots; PITR ≤ 15 min; RPO ≤ 15 min; RTO ≤ 4 h"],
        ],
    )


def section_security(document: Document) -> None:
    add_heading(document, "7. Security & Compliance")
    add_paragraphs(
        document,
        [
            "The solution follows a defense-in-depth strategy and industry standards.",
        ],
    )
    add_bullets(
        document,
        [
            "Authentication & Authorization: OAuth 2.0/OIDC for back-office; secure session for POS; RBAC with least privilege",
            "Password policy and MFA for administrative roles",
            "Data Protection: TLS 1.2+ (pref. TLS 1.3); AES-256 at rest; field-level encryption for sensitive PII",
            "Secret management via cloud KMS/Parameter Store; no secrets in code",
            "PCI DSS alignment: no cardholder data stored/processed by our servers; use P2PE terminals and tokenization",
            "Privacy (GDPR/CCPA): consent tracking, data subject requests, retention policies",
            "Secure coding (OWASP ASVS/Top 10), code scanning, SCA, dependency pinning",
            "Audit logging for authentication, configuration changes, overrides, refunds",
            "Device hardening guidance for store tablets; MDM support optional",
        ],
    )


def section_tech_stack(document: Document) -> None:
    add_heading(document, "8. Technology Stack Recommendation")
    add_paragraphs(
        document,
        [
            "The POS will primarily run on tablets. We recommend a Progressive Web App (PWA) built with React and TypeScript for the Sales interface, leveraging offline storage and background sync. For the Back Office, a responsive React web app shares components and design system.",
        ],
    )
    add_bullets(
        document,
        [
            "Frontend: React + TypeScript (PWA), service worker, IndexedDB (Dexie), state via Redux Toolkit or Zustand; UI via MUI or Tailwind",
            "Hardware access: WebHID/WebUSB/WebBluetooth where supported; optional Capacitor build for broader device support",
            "Backend: Node.js 20 + NestJS; PostgreSQL 15; Prisma ORM; Redis for caching/queues",
            "Realtime/sync: WebSockets; background job processing",
            "Payments: Stripe Terminal (recommended) with tokenization; pluggable provider interface",
            "Cloud: AWS (ALB + ECS/EKS or Fargate), RDS Postgres, S3, CloudFront, Secrets Manager, CloudWatch",
            "CI/CD: GitHub Actions; IaC via Terraform; containerized with Docker",
        ],
    )


def section_data_model(document: Document) -> None:
    add_heading(document, "9. High-Level Data Model (Selected Entities)")
    add_table(
        document,
        headers=["Entity", "Purpose"],
        rows=[
            ["Store", "Physical location with configuration and tax context"],
            ["Register", "Checkout terminal instance at a store"],
            ["User", "System user with role assignments and permissions"],
            ["Customer", "Customer profile, contact details, loyalty data"],
            ["Product", "Base product with brand, category"],
            ["Variant", "Size, color variants; links to SKU"],
            ["SKU", "Sellable unit with barcode and price"],
            ["Inventory", "Stock levels by store/SKU; adjustments and movements"],
            ["Order", "Sales transaction header"],
            ["OrderLine", "Line items referencing SKUs and quantities"],
            ["Payment", "Tender details and tokenized references"],
            ["Promotion", "Discount rules and eligibility"],
            ["TaxRule", "Tax rates and application rules by jurisdiction"],
            ["Supplier", "Vendor information for purchasing"],
            ["PurchaseOrder", "Stock receiving and reconciliation"],
        ],
    )


def section_hardware(document: Document) -> None:
    add_heading(document, "10. Hardware & Peripheral Support")
    add_bullets(
        document,
        [
            "Barcode scanners: Bluetooth LE and USB HID",
            "Receipt printers: ESC/POS network printers (Ethernet/Wi‑Fi); supported models list to be finalized",
            "Cash drawers: via printer kick port",
            "Payment terminals: Stripe Terminal (WisePad 3/BBPOS), Adyen, or Square (pluggable)",
            "Tablet OS: Recent iPadOS and Android versions; Chrome/Safari browsers",
        ],
    )


def section_reports(document: Document) -> None:
    add_heading(document, "11. Reporting & Analytics")
    add_bullets(
        document,
        [
            "Daily sales summary and Z-reports",
            "Product/category performance and sell-through",
            "Cashier performance and overrides/discounts report",
            "Inventory valuation and movement",
            "Tax liability reports",
            "Exports (CSV) and scheduled email delivery",
        ],
    )


def section_devops(document: Document) -> None:
    add_heading(document, "12. Environments, Deployment & Observability")
    add_bullets(
        document,
        [
            "Environments: Dev, Staging, Production with separate resources",
            "CI/CD: automated build/test/lint, vulnerability scans, infrastructure as code",
            "Observability: structured logs, metrics (latency, error rates), tracing; alerts on SLO breaches",
            "Backups: daily snapshots; PITR; tested restore procedures",
        ],
    )


def section_acceptance(document: Document) -> None:
    add_heading(document, "13. Acceptance Criteria (Sample)")
    add_numbered(
        document,
        [
            "As a Sales Associate, I can scan a shoe SKU, select size, and complete a payment with a printed receipt in under 60 seconds.",
            "When the internet is unavailable, I can complete cash sales and the system syncs automatically when connectivity returns without data loss.",
            "Managers can configure a time-bound promotion in Back Office that is applied at POS within 60 seconds of publish.",
            "Inventory received in Back Office updates on-hand counts and is visible at POS within 60 seconds.",
            "All sensitive traffic is encrypted in transit and at rest; administrator actions appear in audit logs within 1 minute.",
        ],
    )


def section_plan(document: Document) -> None:
    add_heading(document, "14. Project Plan & Milestones (Indicative)")
    add_bullets(
        document,
        [
            "Discovery & Design (2–3 weeks): user flows, hardware validation, visual design",
            "MVP Build (6–8 weeks): core POS, catalog, inventory, payments, basic reports",
            "Pilot (2 weeks): limited stores, feedback, bug fixes",
            "Hardening & Training (1–2 weeks): performance tuning, documentation, staff training",
            "General Availability (GA): phased rollout, monitoring, support SLAs",
        ],
    )


def section_assumptions_risks(document: Document) -> None:
    add_heading(document, "15. Assumptions & Risks")
    document.add_paragraph("Assumptions:")
    add_bullets(
        document,
        [
            "Payment provider (e.g., Stripe) account and approved hardware are available",
            "Printer and scanner models selected from a supported list",
            "Tax rules and pricing policies are provided and maintained by client",
        ],
    )
    document.add_paragraph("Risks & Mitigations:")
    add_bullets(
        document,
        [
            "Hardware variability: validate early; maintain compatibility matrix",
            "Offline conflicts: implement clear conflict resolution and operator guidance",
            "Network constraints in-store: support resilient sync and compressed payloads",
        ],
    )


def create_pos_requirements_document(output_path: str) -> None:
    document = Document()
    add_title_page(document)

    section_overview(document)
    section_scope(document)
    section_stakeholders(document)
    section_interfaces(document)
    section_functional_requirements(document)
    section_non_functional(document)
    section_security(document)
    section_tech_stack(document)
    section_data_model(document)
    section_hardware(document)
    section_reports(document)
    section_devops(document)
    section_acceptance(document)
    section_plan(document)
    section_assumptions_risks(document)

    document.save(output_path)


if __name__ == "__main__":
    OUTPUT = "/workspace/POS_System_Requirements.docx"
    create_pos_requirements_document(OUTPUT)
    print(f"Generated: {OUTPUT}")

